"""DOCX 文档解析器 — 输出 Markdown 格式（保留标题层级、列表、表格结构）"""

import io
from typing import Any, Dict, List

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from app.services.parser.base import BaseParser


class DOCXParser(BaseParser):
    """使用 python-docx 提取 DOCX 内容，输出 Markdown 格式"""

    def parse(self, data: bytes, filename: str) -> Dict[str, Any]:
        doc = Document(io.BytesIO(data))
        parts: List[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # 根据样式推断 Markdown 标题层级
            style_name = (para.style.name or "").lower()
            if "heading 1" in style_name:
                parts.append(f"# {text}")
            elif "heading 2" in style_name:
                parts.append(f"## {text}")
            elif "heading 3" in style_name:
                parts.append(f"### {text}")
            elif "heading 4" in style_name:
                parts.append(f"#### {text}")
            elif "list" in style_name:
                parts.append(f"- {text}")
            else:
                # 加粗段落
                if para.runs and all(r.bold for r in para.runs if r.text.strip()):
                    parts.append(f"**{text}**")
                else:
                    parts.append(text)

        # 提取表格为 Markdown 表格
        for table in doc.tables:
            table_lines = _table_to_markdown(table)
            if table_lines:
                parts.append("")
                parts.append(table_lines)

        # 元数据
        metadata: Dict[str, Any] = {"parse_method": "python-docx", "output_format": "markdown"}
        if doc.core_properties.title:
            metadata["title"] = doc.core_properties.title
        if doc.core_properties.author:
            metadata["author"] = doc.core_properties.author

        return {
            "text": "\n\n".join(parts),
            "pages": 0,
            "metadata": metadata,
        }


def _table_to_markdown(table) -> str:
    """将 DOCX 表格转为 Markdown 表格"""
    rows_data: List[List[str]] = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows_data.append(cells)

    if not rows_data:
        return ""

    # 第一行作为表头
    header = rows_data[0]
    separator = ["---"] * len(header)
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(separator) + " |",
    ]
    for row in rows_data[1:]:
        # 补齐列数
        while len(row) < len(header):
            row.append("")
        lines.append("| " + " | ".join(row[:len(header)]) + " |")

    return "\n".join(lines)
